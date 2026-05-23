"""
Resume Export — Markdown to PDF/Word conversion.

Converts the Markdown resume output from resume_rewriter
into styled PDF (.pdf) or Word (.docx) documents.
"""

import os
import re
from datetime import datetime


def export_resume(resume_text: str, format: str = "pdf", output_dir: str = None) -> str:
    """
    Export Markdown resume to PDF or Word document.

    Args:
        resume_text: Markdown-formatted resume text
        format: "pdf" or "docx"
        output_dir: Output directory (defaults to ./outputs/)

    Returns:
        File path to the generated document.
    """
    output_dir = output_dir or os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if format == "pdf":
        return _export_pdf(resume_text, output_dir, timestamp)
    elif format == "docx":
        return _export_docx(resume_text, output_dir, timestamp)
    else:
        raise ValueError(f"Unsupported format: {format}")


def _parse_markdown_sections(text: str) -> list:
    """Parse Markdown into structured sections for document generation."""
    sections = []
    current_section = {"type": "text", "lines": []}

    for line in text.split("\n"):
        stripped = line.strip()

        if stripped.startswith("# ") and not stripped.startswith("## "):
            if current_section["lines"]:
                sections.append(current_section)
            current_section = {"type": "h1", "text": stripped[2:], "lines": []}
        elif stripped.startswith("## "):
            if current_section["lines"] or "text" in current_section:
                sections.append(current_section)
            current_section = {"type": "h2", "text": stripped[3:], "lines": []}
        elif stripped.startswith("### "):
            if current_section["lines"] or "text" in current_section:
                sections.append(current_section)
            current_section = {"type": "h3", "text": stripped[4:], "lines": []}
        elif stripped.startswith("- ") or stripped.startswith("* "):
            current_section.setdefault("bullets", []).append(stripped[2:])
        elif stripped.startswith("|"):
            current_section.setdefault("table", []).append(
                [cell.strip() for cell in stripped.split("|") if cell.strip()]
            )
        elif stripped.startswith("> "):
            current_section.setdefault("quotes", []).append(stripped[2:])
        elif stripped.startswith("**") and "**" in stripped[2:]:
            current_section["lines"].append(stripped)
        elif stripped:
            current_section["lines"].append(stripped)

    if current_section["lines"] or "text" in current_section:
        sections.append(current_section)

    return sections


def _export_pdf(text: str, output_dir: str, timestamp: str) -> str:
    """Export resume as styled PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    except ImportError:
        # Fallback: save as plain text with .pdf.txt extension
        fallback_path = os.path.join(output_dir, f"resume_{timestamp}.txt")
        with open(fallback_path, "w", encoding="utf-8") as f:
            f.write(text)
        return fallback_path + " (reportlab not installed — saved as text)"

    output_path = os.path.join(output_dir, f"resume_{timestamp}.pdf")
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=20*mm, leftMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )

    styles = getSampleStyleSheet()
    accent = HexColor("#3d3dff")
    dark = HexColor("#1a1a20")
    body_color = HexColor("#e8e8ed")

    title_style = ParagraphStyle(
        "ResumeTitle", parent=styles["Heading1"],
        fontSize=18, textColor=accent, spaceAfter=6, spaceBefore=0,
    )
    h2_style = ParagraphStyle(
        "ResumeH2", parent=styles["Heading2"],
        fontSize=14, textColor=accent, spaceAfter=4, spaceBefore=12,
    )
    body_style = ParagraphStyle(
        "ResumeBody", parent=styles["Normal"],
        fontSize=10, textColor=body_color, leading=16,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet", parent=body_style,
        leftIndent=12, bulletIndent=4,
    )
    quote_style = ParagraphStyle(
        "ResumeQuote", parent=body_style,
        leftIndent=16, textColor=HexColor("#9e9eaa"), fontName="Helvetica-Oblique",
    )

    story = []

    sections = _parse_markdown_sections(text)
    for section in sections:
        stype = section.get("type", "text")

        if stype == "h1":
            story.append(Paragraph(_sanitize(section.get("text", "")), title_style))
            story.append(HRFlowable(width="100%", thickness=1, color=accent))
            story.append(Spacer(1, 4*mm))

        elif stype == "h2":
            story.append(Paragraph(_sanitize(section.get("text", "")), h2_style))

        elif stype == "h3":
            story.append(Paragraph(_sanitize(section.get("text", "")), styles["Heading3"]))

        # Body lines
        for line in section.get("lines", []):
            cleaned = _sanitize(line)
            # Bold text
            cleaned = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', cleaned)
            story.append(Paragraph(cleaned, body_style))

        # Bullets
        for bullet in section.get("bullets", []):
            cleaned = _sanitize(bullet)
            cleaned = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', cleaned)
            story.append(Paragraph(f"• {cleaned}", bullet_style))

        # Quotes
        for quote in section.get("quotes", []):
            story.append(Paragraph(_sanitize(quote), quote_style))

        # Tables (minimal rendering)
        table_lines = section.get("table", [])
        if table_lines:
            for row in table_lines:
                story.append(Paragraph(_sanitize(" | ".join(row)), body_style))

        if stype != "h1":
            story.append(Spacer(1, 2*mm))

    doc.build(story)
    return output_path


def _export_docx(text: str, output_dir: str, timestamp: str) -> str:
    """Export resume as styled Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        fallback_path = os.path.join(output_dir, f"resume_{timestamp}.txt")
        with open(fallback_path, "w", encoding="utf-8") as f:
            f.write(text)
        return fallback_path + " (python-docx not installed — saved as text)"

    output_path = os.path.join(output_dir, f"resume_{timestamp}.docx")
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    sections = _parse_markdown_sections(text)
    for section in sections:
        stype = section.get("type", "text")

        if stype == "h1":
            p = doc.add_paragraph()
            run = p.add_run(_sanitize(section.get("text", "")))
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x3d, 0x3d, 0xff)
            doc.add_paragraph()  # spacer

        elif stype == "h2":
            p = doc.add_paragraph()
            run = p.add_run(_sanitize(section.get("text", "")))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x3d, 0x3d, 0xff)

        elif stype == "h3":
            p = doc.add_paragraph()
            run = p.add_run(_sanitize(section.get("text", "")))
            run.bold = True
            run.font.size = Pt(12)

        for line in section.get("lines", []):
            p = doc.add_paragraph()
            _add_formatted_run(p, _sanitize(line))

        for bullet in section.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            _add_formatted_run(p, _sanitize(bullet))

        for quote in section.get("quotes", []):
            p = doc.add_paragraph()
            run = p.add_run(_sanitize(quote))
            run.italic = True
            run.font.size = Pt(10)

    doc.save(output_path)
    return output_path


def _add_formatted_run(paragraph, text: str):
    """Add text with **bold** support to a paragraph."""
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _sanitize(text: str) -> str:
    """Remove problematic characters for document rendering."""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Replace markdown links [text](url) with just text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Strip markdown formatting characters that don't translate
    text = text.replace("`", "")
    return text.strip()
